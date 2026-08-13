from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\Pooja\OneDrive\Desktop\My SDE Resume.docx")
OUTPUT = Path(r"C:\Users\Pooja\OneDrive\Desktop\product-data-management-platform\My SDE Resume - formatted.docx")


def set_no_proof(run):
    """Prevent Word spelling/grammar proofing marks without changing text."""
    rpr = run._r.get_or_add_rPr()
    if rpr.find(qn("w:noProof")) is None:
        rpr.append(OxmlElement("w:noProof"))


def set_bottom_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:pBdr"))
    if old is not None:
        ppr.remove(old)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    ppr.append(borders)


def format_paragraph(p, size=8.5, bold=False, italic=False, align=None,
                     before=0, after=0, left=0, first=0, line=1.0):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.left_indent = Inches(left) if left else None
    fmt.first_line_indent = Inches(first) if first else None
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        set_no_proof(run)


def replace_with_runs(paragraph, parts, size=8.5, align=None, before=0, after=0):
    """Rebuild runs only to apply local emphasis; concatenated text stays identical."""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    format_paragraph(paragraph, size=size, align=align, before=before, after=after)
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        set_no_proof(run)


def set_right_tab(paragraph, inches=7.6):
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(inches), WD_TAB_ALIGNMENT.RIGHT)


def split_on_wide_space(text):
    # Existing documents use large runs of spaces for a right-side date/location.
    import re
    pieces = re.split(r" {5,}", text.strip(), maxsplit=1)
    return "\t".join(pieces) if len(pieces) == 2 else text


def main():
    doc = Document(SOURCE)
    section = doc.sections[0]
    section.top_margin = Inches(0.28)
    section.bottom_margin = Inches(0.28)
    section.left_margin = Inches(0.36)
    section.right_margin = Inches(0.36)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)

    # Keep every existing paragraph and every character; only apply presentation.
    headings = {"EXPERIENCE", "TECHNICAL SKILLS \t", "PROJECTS ", "EDUCATION \t", "ACHIEVEMENTS AND CERTIFICATIONS"}
    bullet_indices = {6, 7, 8, 9, 10, 11, 24, 25, 26, 27, 34, 35, 36}
    employer_indices = {4, 30}
    role_indices = {5, 31}

    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        if idx == 0:
            format_paragraph(p, size=18, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        elif idx in {1, 2}:
            format_paragraph(p, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        elif text in headings:
            format_paragraph(p, size=9.5, bold=True, before=2, after=1)
            set_bottom_border(p)
        elif idx in employer_indices:
            p.text = split_on_wide_space(text)
            format_paragraph(p, size=8.7, bold=True, before=1, after=0)
            set_right_tab(p)
        elif idx in role_indices:
            p.text = split_on_wide_space(text)
            format_paragraph(p, size=8.3, italic=True, after=0)
            set_right_tab(p)
        elif idx == 23:
            # Mirrors the reference: bold project title followed by italic tools.
            prefix = text[:len(text) - len(text.lstrip())]
            clean = text.strip()
            title, tools = clean.split(" | ", 1)
            replace_with_runs(p, [(prefix + title, True, False), (" | ", False, False), (tools, False, True)], size=8.6, before=1)
        elif idx in bullet_indices:
            format_paragraph(p, size=8.1, left=0.16, first=-0.13, after=0)
        elif not text.strip():
            format_paragraph(p, size=2, after=0)
        else:
            format_paragraph(p, size=8.2, after=0)

    # Ensure every run in headers/footers and tables also suppresses proofing.
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for p in container.paragraphs:
                for run in p.runs:
                    set_no_proof(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_no_proof(run)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
